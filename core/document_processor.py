import logging
import os

import docx # python-docx
import nltk # 自然语言工具包 (Natural Language Toolkit)
import uuid # 用于生成唯一的块 ID
from typing import List, Dict, Any, Optional

from config.settings import (
    DEFAULT_PARENT_CHUNK_SIZE, DEFAULT_PARENT_CHUNK_OVERLAP,
    DEFAULT_CHILD_CHUNK_SIZE, DEFAULT_CHILD_CHUNK_OVERLAP,
    SUPPORTED_DOC_EXTENSIONS,
    # DEFAULT_CHUNK_SEPARATOR_REGEX # 如果依赖 NLTK 或段落分割，则不直接使用
)

# 配置日志
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s') # 已在 main 中配置
logger = logging.getLogger(__name__)

# 如果尚未存在，下载 NLTK 的句子分词模型
# 这是一次性设置。在生产环境中，这可能在部署期间处理。
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    logger.info("未找到 NLTK 'punkt' 模型。正在下载...")
    nltk.download('punkt', quiet=True)
    logger.info("'punkt' 模型下载成功。")


class DocumentProcessorError(Exception):
    """DocumentProcessor 错误的自定义异常。"""
    pass

class DocumentProcessor:
    """
    负责处理文档的类：
    - 从各种文件类型（PDF, DOCX, TXT）中提取文本。
    - 将文本分割为父块和子块以用于 RAG。
    """

    def __init__(self,
                 parent_chunk_size: int = DEFAULT_PARENT_CHUNK_SIZE,
                 parent_chunk_overlap: int = DEFAULT_PARENT_CHUNK_OVERLAP,
                 child_chunk_size: int = DEFAULT_CHILD_CHUNK_SIZE,
                 child_chunk_overlap: int = DEFAULT_CHILD_CHUNK_OVERLAP,
                 supported_extensions: List[str] = None):
        """
        初始化 DocumentProcessor。

        Args:
            parent_chunk_size (int): 父块的目标大小（字符数）。
            parent_chunk_overlap (int): 父块之间的重叠部分（字符数）。
            child_chunk_size (int): 子块的目标大小（字符数）。
            child_chunk_overlap (int): 子块之间的重叠部分（字符数）。
            supported_extensions (List[str], optional): 支持的文件扩展名列表。
                                                        默认为 SUPPORTED_DOC_EXTENSIONS。
        """
        self.parent_chunk_size = parent_chunk_size
        self.parent_chunk_overlap = parent_chunk_overlap
        self.child_chunk_size = child_chunk_size
        self.child_chunk_overlap = child_chunk_overlap
        self.supported_extensions = supported_extensions or SUPPORTED_DOC_EXTENSIONS

        # 验证重叠大小（简单验证）
        if self.parent_chunk_overlap >= self.parent_chunk_size and self.parent_chunk_size > 0:
            logger.warning("父块重叠大小 >= 父块大小。正在调整重叠大小。")
            self.parent_chunk_overlap = self.parent_chunk_size // 2 if self.parent_chunk_size > 10 else 0
        if self.child_chunk_overlap >= self.child_chunk_size and self.child_chunk_size > 0:
            logger.warning("子块重叠大小 >= 子块大小。正在调整重叠大小。")
            self.child_chunk_overlap = self.child_chunk_size // 2 if self.child_chunk_size > 10 else 0

        logger.info(f"DocumentProcessor 初始化完成，parent_size={parent_chunk_size}, child_size={child_chunk_size}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        logger.debug(f"正在从 PDF 提取文本 (using pdfplumber): {file_path}")
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # extract_text() usually handles layout better than PyPDF2
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except ImportError:
            logger.error("pdfplumber 未安装。请运行 `pip install pdfplumber`。")
            raise DocumentProcessorError("pdfplumber library not found.")
        except Exception as e:
            logger.error(f"处理 PDF 文件 {file_path} 时出错: {e}")
            raise DocumentProcessorError(f"无法从 PDF {file_path} 提取文本: {e}")

    def _extract_text_from_docx(self, file_path: str) -> str:
        logger.debug(f"正在从 DOCX 提取文本: {file_path}")
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"处理 DOCX 文件 {file_path} 时出错: {e}")
            raise DocumentProcessorError(f"无法从 DOCX {file_path} 提取文本: {e}")

    def _extract_text_from_txt(self, file_path: str) -> str:
        logger.debug(f"正在从 TXT 提取文本: {file_path}")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e: # 捕获潜在的 UnicodeDecodeError 或其他错误
            logger.error(f"处理 TXT 文件 {file_path} 时出错: {e}")
            # 如果 UTF-8 失败，尝试使用回退编码（对 .txt 不常见，但有可能）
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    logger.warning(f"正在使用 'latin-1' 编码重试 TXT 文件 {file_path}。")
                    return file.read()
            except Exception as e_fallback:
                logger.error(f"TXT 文件 {file_path} 的回退编码也失败: {e_fallback}")
                raise DocumentProcessorError(f"无法从 TXT {file_path} 提取文本: {e_fallback}")


    def extract_text_from_file(self, file_path: str) -> str:
        """
        根据文件扩展名提取文件中的文本内容。

        Args:
            file_path (str): 文档文件的路径。

        Returns:
            str: 提取的文本内容。

        Raises:
            DocumentProcessorError: 如果文件类型不支持或处理失败。
            FileNotFoundError: 如果 file_path 不存在。
        """
        logger.info(f"正在尝试从文件提取文本: {file_path}")
        if not os.path.exists(file_path):
            logger.error(f"文件未找到: {file_path}")
            raise FileNotFoundError(f"文件未找到: {file_path}")

        _, extension = os.path.splitext(file_path.lower())

        if extension not in self.supported_extensions:
            msg = f"不支持的文件类型: {extension}。支持的类型有: {self.supported_extensions}"
            logger.error(msg)
            raise DocumentProcessorError(msg)

        extracted_text = ""
        if extension == ".pdf":
            extracted_text = self._extract_text_from_pdf(file_path)
        elif extension == ".docx":
            extracted_text = self._extract_text_from_docx(file_path)
        elif extension == ".txt":
            extracted_text = self._extract_text_from_txt(file_path)

        logger.info(f"成功从 {file_path} 提取了 {len(extracted_text)} 个字符。")
        return extracted_text

    def _recursive_split_text(self, text: str, chunk_size: int, chunk_overlap: int, separators: List[str] = None) -> List[str]:
        """
        递归地分割文本。
        尝试按顺序使用分隔符进行分割。
        如果块仍然太大，则移动到下一个分隔符。
        """
        if separators is None:
            separators = ["\n\n", "\n", " ", ""]

        final_chunks = []
        
        # 获取当前分隔符
        separator = separators[0]
        new_separators = separators[1:]
        
        # 使用当前分隔符分割
        if separator:
            splits = text.split(separator)
        else:
            splits = list(text) # 空分隔符意味着按字符分割

        # 现在的分块可能为空字符串，我们应该过滤掉它们，
        # 还要恢复分隔符（如果不为空），除了最后一个分块
        # 但简单的 split 实际上丢失了分隔符。
        # 对于简单的实现，我们假设分隔符丢失是可以接受的，或者我们可以附加它。
        # 为了保留自然段和句子结构，稍微调整一下：
        # 如果分隔符是标点符号，通常希望保留它。
        # 这里为了简单起见，且通常RAG中分隔符丢失影响不大（除了句子结尾），
        # 我们可以采用 langchain 的 RecursiveCharacterTextSplitter 的策略。
        # 但为了更精准的中文处理，我们可以在分割后将分隔符加回前一个块（如果是句子结尾）。
        
        # 简化版逻辑：先分割，然后尝试合并小块
        good_splits = []
        for s in splits:
             if s.strip():
                 good_splits.append(s)

        current_chunk = ""
        
        for split in good_splits:
            # 尝试将当前 split 添加到 current_chunk
            # 注意：这里我们丢失了分隔符，对于段落 \n\n 可能还好，但对于句子 。 可能会导致句子粘连。
            # 改进：在 split 后面加回 separator (如果是有效字符)
            if separator and separator not in ["\n\n", "\n"]: 
                 split += separator 

            if len(current_chunk) + len(split) + (len(separator) if separator in ["\n\n", "\n"] else 0) <= chunk_size:
                # 能够放入当前块
                 if current_chunk:
                     # 只有当是用换行符分割时，才在合并时加回换行符，或者加空格
                     if separator == "\n\n":
                         current_chunk += "\n\n" + split
                     elif separator == "\n":
                          current_chunk += "\n" + split
                     else:
                          current_chunk += split 
                 else:
                     current_chunk = split
            else:
                # 放入会导致超限
                # 1. 保存当前块（如果有）
                if current_chunk:
                    final_chunks.append(current_chunk)
                    current_chunk = ""
                
                # 2. 处理当前的 split
                if len(split) > chunk_size and new_separators:
                    # 如果单个 split 都太大，且还有更细粒度的分隔符，则递归
                    # 注意：如果刚才加了 separator，递归前可能需要去掉？或者就在这递归处理
                    # 为简单起见，我们对原始 split (不带补回的 separator) 进行递归，或者如果补了就带着。
                    # 如果 split 本身就很大，递归调用
                    sub_chunks = self._recursive_split_text(split, chunk_size, chunk_overlap, new_separators)
                    final_chunks.extend(sub_chunks)
                else:
                    # 没有更多分隔符了，或者 split 小于块大小（虽然不能合并到前一个）
                    # 或者是强制切分（如果没有分隔符了）
                    if len(split) > chunk_size:
                        # 强制切分
                        final_chunks.extend(self._split_into_fixed_size_chunks(split, chunk_size, chunk_overlap))
                    else:
                        current_chunk = split
        
        if current_chunk:
            final_chunks.append(current_chunk)

        return final_chunks

    def split_text_into_parent_child_chunks(self,
                                            full_text: str,
                                            source_document_name: str 
                                           ) -> List[Dict[str, Any]]:
        """
        将文本分割为父块，并将每个父块分割为子块。
        支持中文优先的递归分割。
        """
        if not full_text:
            return []

        doc_name_for_id = os.path.basename(source_document_name)
        structured_chunks = []

        # 1. 拆分为父块 (优先按自然段)
        # 分隔符优先级：双换行(段落) -> 单换行 -> 中文句号 -> 英文句号
        parent_separators = ["\n\n", "\n", "。", "！", "？", ".", "!", "?"]
        
        parent_texts = self._recursive_split_text(
            full_text, 
            self.parent_chunk_size, 
            self.parent_chunk_overlap,
            separators=parent_separators
        )

        logger.info(f"文档 '{source_document_name}' 被拆分为 {len(parent_texts)} 个父候选块。")

        for i, p_text in enumerate(parent_texts):
            if not p_text.strip():
                continue

            parent_id = f"{doc_name_for_id}-p{i+1}"
            parent_chunk_data = {
                "parent_id": parent_id,
                "parent_text": p_text,
                "source_document_name": source_document_name,
                "children": []
            }

            # 2. 将每个父块拆分为子块 (优先按句子)
            # 分隔符优先级：中文标点 -> 英文标点 -> 空格 -> 字符
            child_separators = ["。", "！", "？", "；", "!", "?", ";", "\n", " ", ""]
            
            child_texts = self._recursive_split_text(
                p_text,
                self.child_chunk_size, 
                self.child_chunk_overlap,
                separators=child_separators
            )

            for j, c_text in enumerate(child_texts):
                if not c_text.strip():
                    continue
                child_id = f"{parent_id}-c{j+1}"
                parent_chunk_data["children"].append({
                    "child_id": child_id,
                    "child_text": c_text,
                    "parent_id": parent_id
                })

            if parent_chunk_data["children"]:
                structured_chunks.append(parent_chunk_data)
            else:
                logger.warning(f"父块 '{parent_id}' 未产生有效的子块。跳过此父块。")

        logger.info(f"文档 '{source_document_name}' 已处理为 {len(structured_chunks)} 个带有子块的父块。")
        return structured_chunks


if __name__ == '__main__':
    print("DocumentProcessor 扩展示例")
    logging.basicConfig(level=logging.DEBUG) # 设置为 DEBUG 以从此示例获得更详细的输出

    # 使用默认的父/子大小进行初始化
    processor = DocumentProcessor()

    # --- 测试文本提取 ---
    # 创建用于测试的虚拟文件（在真实场景中，这些文件是存在的）
    dummy_dir = "temp_docs_for_dp_test"
    if not os.path.exists(dummy_dir):
        os.makedirs(dummy_dir)

    dummy_pdf_path = os.path.join(dummy_dir, "test.pdf")
    dummy_docx_path = os.path.join(dummy_dir, "test.docx")
    dummy_txt_path = os.path.join(dummy_dir, "test.txt")
    dummy_unsupported_path = os.path.join(dummy_dir, "test.xls")

    # 创建一个非常简单的 PDF（PyPDF2 无法轻松地将文本写入新 PDF）
    # 所以这个 PDF 可能会提取空文本或在无效时报错。
    # 为了对 PDF 提取进行健壮性测试，需要一个真实的 PDF。
    # 我们将模拟这部分的文本内容。
    try:
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(210, 297) # A4
        with open(dummy_pdf_path, "wb") as f:
            writer.write(f)
        print(f"创建了虚拟 PDF: {dummy_pdf_path}")
    except Exception as e:
        print(f"无法创建虚拟 PDF: {e}")


    # 创建 DOCX
    try:
        doc = docx.Document()
        doc.add_paragraph("This is a test document for DOCX processing.")
        doc.add_paragraph("It contains multiple paragraphs to test extraction.")
        doc.save(dummy_docx_path)
        print(f"创建了虚拟 DOCX: {dummy_docx_path}")
    except Exception as e:
        print(f"无法创建虚拟 DOCX: {e}")


    # 创建 TXT
    try:
        with open(dummy_txt_path, "w", encoding="utf-8") as f:
            f.write("Hello from the TXT file.\nThis is the second line of the text file.")
        print(f"创建了虚拟 TXT: {dummy_txt_path}")
    except Exception as e:
        print(f"无法创建虚拟 TXT: {e}")

    with open(dummy_unsupported_path, "w") as f: f.write("content") # 虚拟不支持的文件

    file_paths_to_test = {
        "pdf": dummy_pdf_path,
        "docx": dummy_docx_path,
        "txt": dummy_txt_path,
        "unsupported": dummy_unsupported_path
    }

    # 模拟 PDF 文本，因为 PyPDF2 需要一个包含真实文本的 PDF
    MOCK_PDF_TEXT = "This is mocked PDF text. It has sentences. And paragraphs too.\n\nThis is a new paragraph in the mocked PDF."

    for file_type, path in file_paths_to_test.items():
        print(f"\n--- 测试 {file_type.upper()} 文件的提取: {path} ---")
        try:
            if file_type == "pdf" and os.path.exists(path): # 对 PDF 使用模拟
                 text_content = MOCK_PDF_TEXT
                 print(f"对 PDF 使用模拟文本: '{text_content[:100]}...'")
            elif os.path.exists(path):
                text_content = processor.extract_text_from_file(path)
                print(f"提取的文本 ({len(text_content)} 字符): '{text_content[:100]}...'")
            else:
                print(f"文件 {path} 不存在。跳过。")
                continue

            # --- 测试中文内容分块 ---
            # 如果是 txt 且我们想测试中文逻辑，可以覆盖内容
            if file_type == "txt":
                text_content = (
                    "这是一个用于测试中文分块的长文本。第一段包含几个句子。这是第一句。"
                    "这是第二句，稍微长一点，看看能不能正确分割。这是第三句。\n\n"
                    "这是第二段。它应该被识别为一个新的父块。这一段也包含一些句子。"
                    "为了测试递归分割，我们需要一个非常非常长的句子，它甚至可能超过子块的大小限制，"
                    "所以我们需要看看递归分割器是否能够正确地处理这种情况，比如通过逗号或者其他符号来进一步分割，"
                    "或者如果实在不行，就强制切分。希望这个逻辑能工作正常。"
                )
                print(f"覆盖为中文测试文本: '{text_content[:50]}...'")

            # --- 测试父子分块 ---
            if text_content and text_content.strip():
                print(f"\n--- 测试 {file_type.upper()} 内容的父子分块 ---")
                # 对此示例文本使用较小的块大小以查看更多块
                processor_small_chunks = DocumentProcessor(parent_chunk_size=80, parent_chunk_overlap=10,
                                                           child_chunk_size=30, child_chunk_overlap=5)

                parent_child_chunks = processor_small_chunks.split_text_into_parent_child_chunks(
                    text_content, source_document_name=path # 将文件路径作为名称传递
                )

                print(f"生成了 {len(parent_child_chunks)} 个父块。")
                for i, parent_data in enumerate(parent_child_chunks):
                    print(f"  父块 {i+1} (ID: {parent_data['parent_id']}): '{parent_data['parent_text'][:50]}...'")
                    print(f"    源文档名称: {parent_data['source_document_name']}") # 更新键
                    print(f"    包含 {len(parent_data['children'])} 个子块:")
                    for j, child_data in enumerate(parent_data['children']):
                        print(f"      子块 {j+1} (ID: {child_data['child_id']}): '{child_data['child_text'][:40]}...'")
                if not parent_child_chunks:
                     print("未生成父子块（文本对于当前设置可能太短）。")
            else:
                print("由于提取的文本为空，跳过分块。")

        except (FileNotFoundError, DocumentProcessorError) as e:
            print(f"处理 {file_type} 期间出错: {e}")
        except Exception as e:
            print(f"对于 {file_type} 发生意外错误: {e}")
            # import traceback; traceback.print_exc()


    # 清理虚拟文件
    try:
        if os.path.exists(dummy_dir):
            import shutil
            shutil.rmtree(dummy_dir)
            print(f"\n清理了临时目录: {dummy_dir}")
    except Exception as e:
        print(f"清理临时目录时出错: {e}")

    print("\nDocumentProcessor 扩展示例结束。")
